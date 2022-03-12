import docx  
import sys

class WordReader:

    def __init__(self, docx_path=""):
        self.document = docx.Document(docx_path)

    def get_paragraphs(self):
        return self.document.paragraphs
        #for i in docFile.paragraphs:
        #    for j in i.runs:
        #        print(j.font.size/12700)

    def get_styles(self):
        styles =self.document.styles
        paragraph_styles = [
            s for s in styles if s.type == docx.enum.style.WD_STYLE_TYPE.PARAGRAPH
        ]
        return paragraph_styles



if __name__ == '__main__':
    if len(sys.argv) != 2:
          raise Exception("Debe especificar la ruta del word")

    word_path = sys.argv[1]     
    word_reader = WordReader(word_path)
    
    for style in word_reader.get_styles():
        print(style.name)
        print(style.font.name)
        print(style.font.size / 12700)
        print("=======")
    #for par in word_reader.get_paragraphs():
        # print(par.style)
        # for run in par.runs:
        #    print(run.font.size/12700) 
        #    print("---")       
        # print(par.text)